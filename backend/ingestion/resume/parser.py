"""Resume parser for DOCX and PDF files."""

import json
import re
from typing import Dict, Any, Optional, List
from datetime import datetime
import logging
from pathlib import Path

import docx
import PyPDF2

from ...models.profile import Resume
from ...core.storage import storage_manager

logger = logging.getLogger(__name__)


class ResumeParser:
    """Parse resume files into structured JSON."""
    
    def __init__(self):
        self.section_patterns = {
            "summary": ["summary", "professional summary", "about", "profile"],
            "experience": ["experience", "work experience", "employment", "work history"],
            "education": ["education", "academic background", "qualifications"],
            "skills": ["skills", "technical skills", "competencies", "expertise"],
            "projects": ["projects", "personal projects", "key projects"],
            "certifications": ["certifications", "licenses", "credentials"],
            "languages": ["languages", "spoken languages"],
            "publications": ["publications", "papers", "research"],
            "awards": ["awards", "honors", "achievements"]
        }
    
    async def parse_file(self, file_path: Path, file_type: str) -> Dict[str, Any]:
        """Parse resume file based on type."""
        if file_type.lower() == "docx":
            return self._parse_docx(file_path)
        elif file_type.lower() == "pdf":
            return self._parse_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX resume."""
        doc = docx.Document(file_path)
        
        # Extract text with formatting
        full_text = []
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name if para.style else None,
                    "runs": [{"text": run.text, "bold": run.bold, "italic": run.italic} 
                            for run in para.runs]
                })
                full_text.append(para.text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return self._structure_resume("\n".join(full_text), paragraphs, tables)
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF resume."""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            full_text = []
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)
            
            return self._structure_resume("\n".join(full_text))
    
    def _structure_resume(self, text: str, paragraphs: Optional[List] = None, 
                          tables: Optional[List] = None) -> Dict[str, Any]:
        """Structure raw text into sections."""
        lines = text.split('\n')
        
        # Extract sections
        sections = self._extract_sections(lines)
        
        # Extract personal info
        personal_info = self._extract_personal_info(text)
        
        # Extract skills
        skills = self._extract_skills(text)
        
        # Extract experience
        experience = self._parse_experience(sections.get("experience", []))
        
        # Extract education
        education = self._parse_education(sections.get("education", []))
        
        # Extract projects
        projects = self._parse_projects(sections.get("projects", []))
        
        return {
            "personal": personal_info,
            "summary": sections.get("summary", [""])[0] if sections.get("summary") else "",
            "skills": skills,
            "experience": experience,
            "education": education,
            "projects": projects,
            "certifications": self._parse_certifications(sections.get("certifications", [])),
            "languages": self._parse_languages(sections.get("languages", [])),
            "metadata": {
                "has_tables": bool(tables),
                "has_paragraphs": bool(paragraphs),
                "line_count": len(lines),
                "word_count": len(text.split()),
                "sections_found": list(sections.keys())
            }
        }
    
    def _extract_sections(self, lines: List[str]) -> Dict[str, List[str]]:
        """Extract sections based on headers."""
        sections = {}
        current_section = "header"
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            is_header = False
            matched_section = None
            
            # Common header patterns
            if line.isupper() and len(line) > 3:
                is_header = True
                matched_section = line.lower()
            elif line.endswith(':') and len(line) < 30:
                is_header = True
                matched_section = line[:-1].lower()
            else:
                # Check against known section names
                line_lower = line.lower()
                for section, patterns in self.section_patterns.items():
                    if any(pattern in line_lower for pattern in patterns):
                        if len(line) < 50:  # Likely a header
                            is_header = True
                            matched_section = section
                            break
            
            if is_header and matched_section:
                # Save previous section
                if current_content:
                    sections[current_section] = current_content
                
                # Start new section
                current_section = matched_section
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = current_content
        
        return sections
    
    def _extract_personal_info(self, text: str) -> Dict[str, Any]:
        """Extract personal information from text."""
        info = {
            "name": None,
            "email": None,
            "phone": None,
            "location": None,
            "linkedin": None,
            "github": None,
            "website": None
        }
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            info["email"] = emails[0]
        
        # Phone pattern (simple)
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            info["phone"] = phones[0]
        
        # LinkedIn URL
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            info["linkedin"] = linkedin.group()
        
        # GitHub URL
        github_pattern = r'github\.com/[\w-]+'
        github = re.search(github_pattern, text, re.IGNORECASE)
        if github:
            info["github"] = github.group()
        
        return info
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text."""
        # Common skill keywords (could be expanded)
        skill_keywords = [
            "python", "javascript", "typescript", "java", "c++", "c#", "ruby", "php",
            "react", "angular", "vue", "node.js", "express", "django", "flask", "fastapi",
            "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins",
            "sql", "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
            "git", "github", "gitlab", "jira", "confluence",
            "agile", "scrum", "kanban", "tdd", "ci/cd"
        ]
        
        found_skills = []
        text_lower = text.lower()
        
        for skill in skill_keywords:
            if skill in text_lower:
                found_skills.append(skill)
        
        # Also look for comma-separated lists
        skill_sections = re.findall(r'(?:skills|technologies|expertise)[:\s]+(.*?)(?:\n\n|\Z)', 
                                     text_lower, re.IGNORECASE | re.DOTALL)
        
        for section in skill_sections:
            # Split by common separators
            items = re.split(r'[,|•\n]', section)
            for item in items:
                item = item.strip()
                if item and len(item) < 50:  # Avoid long sentences
                    found_skills.append(item)
        
        return list(set(found_skills))  # Deduplicate
    
    def _parse_experience(self, lines: List[str]) -> List[Dict[str, Any]]:
        """Parse experience section."""
        experiences = []
        current_exp = {}
        
        for line in lines:
            # Look for date patterns
            date_pattern = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{4}\s*[-–]\s*(?:present|\d{4})'
            if re.search(date_pattern, line, re.IGNORECASE):
                if current_exp:
                    experiences.append(current_exp)
                
                # Try to extract company and title
                parts = line.split(' at ') if ' at ' in line else line.split(' - ')
                if len(parts) > 1:
                    title = parts[0].strip()
                    company = parts[1].strip()
                else:
                    title = line
                    company = None
                
                current_exp = {
                    "title": title,
                    "company": company,
                    "dates": line,
                    "description": []
                }
            elif current_exp:
                if line and not line.startswith(('•', '-', '*')):
                    current_exp["description"].append(line)
                elif line.startswith(('•', '-', '*')):
                    # Bullet point
                    current_exp["description"].append(line[1:].strip())
        
        if current_exp:
            experiences.append(current_exp)
        
        return experiences
    
    def _parse_education(self, lines: List[str]) -> List[Dict[str, Any]]:
        """Parse education section."""
        education = []
        
        for line in lines:
            # Look for degree indicators
            if any(deg in line.lower() for deg in ['bachelor', 'master', 'phd', 'b.s.', 'm.s.', 'b.a.', 'm.a.']):
                education.append({
                    "degree": line,
                    "institution": None,
                    "dates": None
                })
        
        return education
    
    def _parse_projects(self, lines: List[str]) -> List[Dict[str, Any]]:
        """Parse projects section."""
        projects = []
        current_project = {}
        
        for line in lines:
            if line and not line.startswith(('•', '-', '*')) and not current_project:
                # New project title
                current_project = {
                    "name": line,
                    "description": []
                }
            elif current_project:
                if line.startswith(('•', '-', '*')):
                    # Project detail
                    current_project["description"].append(line[1:].strip())
                elif not line:
                    # Empty line - end of project
                    if current_project:
                        projects.append(current_project)
                        current_project = {}
        
        if current_project:
            projects.append(current_project)
        
        return projects
    
    def _parse_certifications(self, lines: List[str]) -> List[str]:
        """Parse certifications section."""
        certs = []
        for line in lines:
            if line and not line.startswith(('•', '-', '*')):
                certs.append(line)
        return certs
    
    def _parse_languages(self, lines: List[str]) -> List[Dict[str, str]]:
        """Parse languages section."""
        languages = []
        for line in lines:
            if line and not line.startswith(('•', '-', '*')):
                # Check for proficiency indicators
                parts = line.split(' - ')
                if len(parts) > 1:
                    languages.append({
                        "language": parts[0].strip(),
                        "proficiency": parts[1].strip()
                    })
                else:
                    languages.append({
                        "language": line,
                        "proficiency": None
                    })
        return languages


class ResumeStorage:
    """Handle resume storage and retrieval."""
    
    def __init__(self):
        self.parser = ResumeParser()
    
    async def store_canonical_resume(self, file_path: Path, file_type: str, profile_id: int) -> int:
        """Parse and store canonical resume."""
        # Parse resume
        parsed_data = await self.parser.parse_file(file_path, file_type)
        
        # Upload to storage
        storage_key = f"resumes/canonical/{profile_id}/resume.{file_type}"
        with open(file_path, 'rb') as f:
            file_url = await storage_manager.upload_file(
                key=storage_key,
                file_data=f.read(),
                content_type=f"application/{file_type}"
            )
        
        # Create resume record
        async with db_manager.session() as session:
            from ...models.profile import Resume
            
            resume = Resume(
                profile_id=profile_id,
                version="canonical",
                name="Canonical Resume",
                is_canonical=True,
                file_path=storage_key,
                file_type=file_type,
                file_size=file_path.stat().st_size,
                content_json=parsed_data
            )
            
            session.add(resume)
            await session.commit()
            await session.refresh(resume)
            
            return resume.id
    
    async def get_canonical_resume(self, profile_id: int) -> Optional[Dict[str, Any]]:
        """Retrieve canonical resume for a profile."""
        async with db_manager.session() as session:
            from sqlalchemy import select
            from ...models.profile import Resume
            
            stmt = select(Resume).where(
                Resume.profile_id == profile_id,
                Resume.is_canonical == True
            )
            result = await session.execute(stmt)
            resume = result.scalar_one_or_none()
            
            if resume:
                return resume.content_json
            return None